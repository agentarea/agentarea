"""End-to-end docx-skill integration test.

Proves the AgentArea skill system + per-workflow sandbox actually work together:

  1. Upload a docx-generator skill (zipped SKILL.md + make_docx.py).
  2. Create an agent with skill_ids=[<docx_skill>] and no shell tool.
  3. Submit a sync task that asks for a commercial-offer DOCX.
  4. Poll the workflow events until terminal.
  5. Assert the agent called activate_skill + run_skill_script (and not bash).
  6. Pull the produced .docx out of the sandbox via the same workflow_id and
     parse it with python-docx to confirm it's a real Office Open XML file
     containing the requested customer name.

Run with:
    cd agentarea-platform
    uv run pytest -m integration tests/e2e/test_docx_skill_integration.py -v -s

Required: full docker-compose stack up (make up-dev). The mcp-manager must be
reachable at http://localhost:7999 (default port mapping).
"""

from __future__ import annotations

import base64
import hashlib
import io
import os
import time
import uuid
import zipfile
from pathlib import Path

import httpx
import pytest

from tests.e2e.api.conftest import _psql, wait_for_workflow
from tests.e2e.api.test_skill_mcp_orchestration import _ensure_artifacts_bucket

API_URL = os.environ.get("API_URL", "http://localhost:8000")
MCP_MANAGER_URL = os.environ.get("MCP_MANAGER_URL", "http://localhost:7999")

# OpenRouter / Kimi K2 credentials must come from the environment.
# DO NOT hardcode keys here — this file is checked into version control.
OPENROUTER_API_KEY = os.environ.get("DOCX_E2E_OPENROUTER_KEY", "")
OPENROUTER_ENDPOINT = os.environ.get("DOCX_E2E_OPENROUTER_ENDPOINT", "https://openrouter.ai/api/v1")
OPENROUTER_MODEL = os.environ.get("DOCX_E2E_OPENROUTER_MODEL", "moonshotai/kimi-k2.6")
OPENROUTER_PROVIDER_KEY = os.environ.get("DOCX_E2E_PROVIDER_KEY", "e2e-openrouter-docx")

pytestmark = pytest.mark.skipif(
    not OPENROUTER_API_KEY,
    reason=(
        "DOCX_E2E_OPENROUTER_KEY is not set. Export your OpenRouter API key "
        "before running this integration test, e.g. "
        "`export DOCX_E2E_OPENROUTER_KEY=sk-or-v1-...`"
    ),
)

ARTIFACT_DIR = Path(
    os.environ.get(
        "DOCX_E2E_ARTIFACT_DIR",
        str(Path(__file__).resolve().parents[3] / ".omc" / "artifacts" / "docx-e2e"),
    )
)


# ---------------------------------------------------------------------------
# Provider/model fixtures (session scoped — match conftest pattern but pinned
# to OpenRouter + Kimi instead of the default OpenAI-compat proxy).
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def openrouter_provider_spec_id() -> str:
    _psql(
        "INSERT INTO provider_specs(id,provider_key,name,provider_type,"
        "is_builtin,workspace_id,created_by) "
        f"VALUES (gen_random_uuid(),'{OPENROUTER_PROVIDER_KEY}',"
        "'OpenRouter (docx-e2e)','openai-compatible',true,'system','system') "
        "ON CONFLICT (provider_key) DO NOTHING;"
    )
    spec_id = _psql(
        f"SELECT id FROM provider_specs WHERE provider_key='{OPENROUTER_PROVIDER_KEY}';"
    )
    assert spec_id, "failed to seed openrouter provider_spec"
    return spec_id


@pytest.fixture(scope="session")
def kimi_model_spec_id(openrouter_provider_spec_id: str) -> str:
    _psql(
        "UPDATE model_specs SET workspace_id='system', created_by='system' "
        f"WHERE provider_spec_id='{openrouter_provider_spec_id}' "
        f"AND model_name='{OPENROUTER_MODEL}';"
    )
    _psql(
        "INSERT INTO model_specs(id,provider_spec_id,model_name,display_name,"
        "context_window,is_active,workspace_id,created_by) VALUES "
        f"(gen_random_uuid(),'{openrouter_provider_spec_id}','{OPENROUTER_MODEL}',"
        f"'{OPENROUTER_MODEL}',128000,true,'system','system') "
        "ON CONFLICT (provider_spec_id, model_name) DO NOTHING;"
    )
    spec_id = _psql(
        f"SELECT id FROM model_specs WHERE provider_spec_id='{openrouter_provider_spec_id}' "
        f"AND model_name='{OPENROUTER_MODEL}';"
    )
    assert spec_id, "failed to seed kimi model_spec"
    return spec_id


@pytest.fixture
def kimi_model(
    alice_client: httpx.Client,
    openrouter_provider_spec_id: str,
    kimi_model_spec_id: str,
) -> str:
    pc = (
        alice_client.post(
            "/v1/provider-configs/",
            json={
                "provider_spec_id": openrouter_provider_spec_id,
                "name": f"docx-e2e-{uuid.uuid4().hex[:6]}",
                "api_key": OPENROUTER_API_KEY,
                "endpoint_url": OPENROUTER_ENDPOINT,
            },
        )
        .raise_for_status()
        .json()
    )

    mi = (
        alice_client.post(
            "/v1/model-instances/",
            json={
                "provider_config_id": pc["id"],
                "model_spec_id": kimi_model_spec_id,
                "name": f"docx-e2e-{uuid.uuid4().hex[:6]}",
            },
        )
        .raise_for_status()
        .json()
    )
    return mi["id"]


# ---------------------------------------------------------------------------
# Skill payload
# ---------------------------------------------------------------------------

SKILL_NAME = "docx-generator"
SKILL_DESCRIPTION = (
    "Generate real Microsoft Word (.docx) files for commercial offers, letters, and reports."
)
OUTPUT_FILENAME = "output.docx"

SKILL_MD = f"""---
name: {SKILL_NAME}
description: {SKILL_DESCRIPTION}
---

# DOCX Generator

Use this skill when the user wants a Microsoft Word (.docx) document
(commercial offer, letter, report, etc.).

## How to use (CRITICAL — read carefully)

The sandbox argv sanitizer REJECTS quotes, braces, parens, and shell
metacharacters. To pass structured data into `make_docx.py`, the script
uses a flat **comma-separated key/value** format. Spaces inside values
are fine; commas and equals signs are reserved.

### Argument format

A single string with this shape:

  `key1=value1,key2=value2,key3=value3`

The script understands these five keys (all optional except `client`):

  - `title`     document title shown at the top
  - `client`    REQUIRED — client / addressee organization
  - `recipient` person being addressed (e.g. John Smith)
  - `intro`     one paragraph of intro text
  - `closing`   closing paragraph

Plain text values only. No semicolons, no quotes, no braces — the
sandbox rejects shell metacharacters in arguments. If the user mentions
pricing or scope, fold it into the `intro` paragraph as prose.

### Tool call

Call `run_skill_script` exactly once with all five keys filled in:

  `run_skill_script(skill_name={SKILL_NAME},
  script_name=make_docx.py,
  args=title=Commercial Offer,client=ACME Corp,recipient=John Smith,intro=...,closing=...)`

### Output

The script prints exactly one line on success:

  `WROTE: path=/workspace/wf-<id>/{OUTPUT_FILENAME}, size=<bytes>,
  sha256=<hex>, b64=<base64>`

Treat that as proof. After you see it, call `completion` with one short
sentence that includes the client name and confirms the .docx was
written.

## Constraints

- Do NOT use bash, shell, or any other tool — only `run_skill_script`.
- Do NOT use JSON / quotes / curly braces in `args`. The sandbox will
  reject them and return exit code 1.
- Keep `args` short — under 1000 characters.
- If the script returns an error, fix the args and retry ONCE.
"""

MAKE_DOCX_PY = '''\
"""Generate a real .docx file from a JSON spec passed as sys.argv[1]."""

from __future__ import annotations

import base64
import hashlib
import json
import os
import subprocess
import sys


def _ensure_python_docx() -> None:
    """Install python-docx into a per-workflow venv inside the workspace.

    The venv lives at /workspace/wf-<id>/.venv — same lifetime as the
    workflow's sandbox state. Two consequences:
      * the cost of `pip install python-docx` is paid once per task and
        amortized across every script call in that workflow;
      * nothing leaks across workflows. The next workflow on the same
        warm-pool pod starts with a clean workspace and re-creates its
        own venv. No shared system-level pip install.
    """
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        pass

    import venv

    venv_dir = os.path.join(os.getcwd(), ".venv")
    py = os.path.join(venv_dir, "bin", "python")
    if not os.path.isdir(venv_dir):
        venv.create(venv_dir, with_pip=True)
        subprocess.run(
            [py, "-m", "pip", "install", "--quiet", "--no-input", "python-docx"],
            check=True,
        )
    # Re-exec under the venv so the import resolves to the installed package.
    os.execv(py, [py, __file__, *sys.argv[1:]])


def _parse_kv_args(raw: str) -> dict:
    """Parse a comma-separated key=value string into a dict.

    Repeated `item` keys collect into a list. Values may contain spaces but
    not commas or equals signs (commas split entries; the first `=` per
    entry splits key from value).
    """
    spec: dict = {"items": []}
    for entry in raw.split(","):
        entry = entry.strip()
        if not entry or "=" not in entry:
            continue
        key, _, value = entry.partition("=")
        key = key.strip()
        value = value.strip()
        if not key:
            continue
        if key == "item":
            parts = [p.strip() for p in value.split(";")]
            while len(parts) < 3:
                parts.append("")
            spec["items"].append(
                {"name": parts[0], "description": parts[1], "price": parts[2]}
            )
        else:
            spec[key] = value
    return spec


def main() -> int:
    if len(sys.argv) < 2:
        print(
            "Error: pass a comma-separated key=value string as argv[1]",
            file=sys.stderr,
        )
        return 2

    raw = sys.argv[1]
    # Accept either the kv-format (default skill contract) or raw JSON when
    # the script is invoked manually for debugging.
    if raw.lstrip().startswith("{"):
        try:
            spec = json.loads(raw)
        except json.JSONDecodeError as e:
            print(f"Error: invalid JSON: {e}", file=sys.stderr)
            return 2
    else:
        spec = _parse_kv_args(raw)

    _ensure_python_docx()
    from docx import Document

    doc = Document()

    title = str(spec.get("title", "Commercial Offer"))
    doc.add_heading(title, level=0)

    client = str(spec.get("client", ""))
    if client:
        doc.add_paragraph().add_run(f"For: {client}").bold = True

    recipient = str(spec.get("recipient", ""))
    if recipient:
        doc.add_paragraph(f"Attn: {recipient}")

    intro = str(spec.get("intro", ""))
    if intro:
        doc.add_paragraph(intro)

    items = spec.get("items") or []
    if items:
        doc.add_heading("Line Items", level=1)
        table = doc.add_table(rows=1 + len(items), cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Description"
        hdr[2].text = "Price"
        for i, item in enumerate(items, start=1):
            row = table.rows[i].cells
            row[0].text = str(item.get("name", ""))
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("price", ""))

    closing = str(spec.get("closing", ""))
    if closing:
        doc.add_paragraph(closing)

    out_name = "output.docx"
    cwd = os.getcwd()
    out_path = os.path.join(cwd, out_name)
    doc.save(out_path)

    with open(out_path, "rb") as f:
        data = f.read()

    sha = hashlib.sha256(data).hexdigest()
    b64 = base64.b64encode(data).decode("ascii")
    # Print on one line so the LLM\'s tool result + the test\'s parser both
    # have an easy fixed-format anchor.
    print(
        f"WROTE: path={out_path}, size={len(data)}, sha256={sha}, b64={b64}",
        flush=True,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
'''


def _build_skill_zip() -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("SKILL.md", SKILL_MD)
        zf.writestr("make_docx.py", MAKE_DOCX_PY)
    return buf.getvalue()


@pytest.fixture
def docx_skill_id(alice_client: httpx.Client) -> str:
    _ensure_artifacts_bucket()
    zip_bytes = _build_skill_zip()
    files = {
        "file": (
            "docx-generator.zip",
            zip_bytes,
            "application/zip",
        )
    }
    # Each fixture run uses a fresh ephemeral Kratos user (workspace), so the
    # bare skill name from SKILL.md frontmatter never collides.
    resp = alice_client.post(
        "/v1/skills/upload",
        files=files,
        timeout=30.0,
    )
    resp.raise_for_status()
    return resp.json()["id"]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _tool_calls_observed(events: list[dict]) -> list[str]:
    """Names of tools the agent invoked (deduped insertion-order list)."""
    names: list[str] = []
    seen: set[str] = set()
    for ev in events:
        if ev["event_type"] not in {"ToolCallStarted", "ToolCallCompleted"}:
            continue
        name = ev.get("metadata", {}).get("tool_name")
        if name and name not in seen:
            seen.add(name)
            names.append(name)
    return names


def _read_docx_from_sandbox(workflow_id: str) -> bytes:
    """Pull /workspace/wf-<workflow_id>/output.docx out via mcp-manager.

    Uses a fresh script invocation that base64-emits the file so we don't
    rely on the agent's tool result being parseable from event metadata.
    """
    payload = {
        "workflow_id": workflow_id,
        "script_name": "read.sh",
        "script_content": (
            "if [ ! -f output.docx ]; then echo MISSING; exit 1; fi; base64 output.docx"
        ),
        "timeout_seconds": 30,
    }
    resp = httpx.post(f"{MCP_MANAGER_URL}/sandbox/execute", json=payload, timeout=40.0)
    resp.raise_for_status()
    body = resp.json()
    if body.get("exit_code") != 0:
        raise AssertionError(
            f"sandbox readback failed: exit={body.get('exit_code')} "
            f"stdout={body.get('stdout')!r} stderr={body.get('stderr')!r}"
        )
    encoded = (body.get("stdout") or "").strip()
    if not encoded or encoded == "MISSING":
        raise AssertionError(f"output.docx not present in workspace: {body!r}")
    # base64 may include newlines; b64decode tolerates them.
    return base64.b64decode(encoded)


# ---------------------------------------------------------------------------
# The actual test
# ---------------------------------------------------------------------------


@pytest.mark.integration
@pytest.mark.slow
def test_docx_skill_end_to_end(
    alice_client: httpx.Client,
    kimi_model: str,
    docx_skill_id: str,
) -> None:
    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)

    instruction = (
        "You are a document-writing assistant.\n"
        "When the user asks for a Word (.docx) document, you MUST:\n"
        f"  1. Call activate_skill(skill_name='{SKILL_NAME}') to load the "
        "full skill instructions.\n"
        "  2. Carefully read the SKILL.md instructions returned. The skill "
        "requires a flat comma-separated key=value string for `args` "
        "(NO JSON, NO quotes, NO braces). Follow that contract exactly.\n"
        f"  3. Call run_skill_script with skill_name='{SKILL_NAME}', "
        "script_name='make_docx.py', and args=<KEY_VALUE_STRING>.\n"
        "  4. After the script prints its WROTE: line, call completion with "
        "a one-sentence summary mentioning the client name.\n"
        "Never use bash, shell, or any other tool. Only the skill tools."
    )

    create_resp = alice_client.post(
        "/v1/agents/",
        json={
            "name": f"docx-tester-{uuid.uuid4().hex[:6]}",
            "description": "docx-skill e2e",
            "instruction": instruction,
            "model_id": kimi_model,
            "agent_type": "stateless",
            "tools": [],
            "skill_ids": [docx_skill_id],
        },
        timeout=15.0,
    )
    create_resp.raise_for_status()
    agent_id = create_resp.json()["id"]

    task_description = (
        "Generate a commercial offer in DOCX format for ACME Corp. "
        "Cover: implementation of a per-workflow sandbox. "
        "Pricing: $0 across the board (this is a demo). "
        "Address it to John Smith, Procurement Director."
    )

    task_resp = alice_client.post(
        f"/v1/agents/{agent_id}/tasks/sync",
        json={"description": task_description},
        timeout=30.0,
    )
    task_resp.raise_for_status()
    task_id = task_resp.json()["id"]

    print(f"\n[docx-e2e] agent_id={agent_id} task_id={task_id}")

    started = time.time()
    # 600s headroom: kimi-k2.6 has been observed taking ~25s/iteration in
    # the worst case, and the workflow may run 10+ iterations if the model
    # second-guesses its argument format on the first call. The actual
    # workflow termination is detected via WorkflowCompleted/WorkflowFailed
    # events emitted in agent_execution_workflow._finalize_execution.
    events = wait_for_workflow(
        alice_client,
        agent_id,
        task_id,
        timeout=600.0,
        poll=2.0,
    )
    elapsed = time.time() - started
    print(f"[docx-e2e] workflow terminated in {elapsed:.1f}s, {len(events)} events")

    # Tool path assertions ---------------------------------------------------
    tool_names = _tool_calls_observed(events)
    print(f"[docx-e2e] tool calls observed: {tool_names}")

    assert "activate_skill" in tool_names, (
        f"expected activate_skill in tool calls, got {tool_names!r}"
    )
    assert "run_skill_script" in tool_names, (
        f"expected run_skill_script in tool calls, got {tool_names!r}"
    )
    assert "bash" not in tool_names, f"agent should not have access to bash, got {tool_names!r}"

    # Workflow terminal state ------------------------------------------------
    terminal = next(
        (
            ev
            for ev in reversed(events)
            if ev["event_type"] in {"WorkflowCompleted", "WorkflowFailed"}
        ),
        None,
    )
    assert terminal is not None, "no terminal event found"
    assert terminal["event_type"] == "WorkflowCompleted", (
        f"workflow did not complete cleanly; terminal={terminal!r}"
    )

    # Confirm the run_skill_script call actually fired and didn't crash.
    skill_script_completions = [
        ev
        for ev in events
        if ev["event_type"] == "ToolCallCompleted"
        and ev.get("metadata", {}).get("tool_name") == "run_skill_script"
    ]
    assert skill_script_completions, "no ToolCallCompleted for run_skill_script"
    last_run = skill_script_completions[-1]["metadata"]
    # exit_code is set on the ad-hoc skill_script TOOL_CALL_COMPLETED payload.
    if "exit_code" in last_run:
        assert last_run["exit_code"] == 0, (
            f"run_skill_script returned non-zero exit code: {last_run!r}"
        )

    # Pull the .docx out of the sandbox -------------------------------------
    workflow_id = f"task-{task_id}"
    docx_bytes = _read_docx_from_sandbox(workflow_id)

    sha = hashlib.sha256(docx_bytes).hexdigest()
    saved_path = ARTIFACT_DIR / f"{task_id}-{sha[:12]}.docx"
    saved_path.write_bytes(docx_bytes)
    print(f"[docx-e2e] saved {len(docx_bytes)} bytes to {saved_path} sha256={sha}")

    # Parse with python-docx -------------------------------------------------
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    paragraphs_text = [p.text for p in doc.paragraphs]
    table_texts: list[str] = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                table_texts.append(cell.text)

    full_text = "\n".join(paragraphs_text + table_texts)
    print(f"[docx-e2e] doc text preview:\n{full_text[:500]}")

    assert paragraphs_text, "docx had no paragraphs at all"
    assert any(p.strip() for p in paragraphs_text), "all paragraphs empty"
    assert "ACME" in full_text or "Acme" in full_text or "acme" in full_text, (
        f"customer 'ACME Corp' missing from docx; got:\n{full_text[:500]}"
    )
